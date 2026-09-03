from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "report-assets-mocthu"
OUTPUT = DOCS / "BAO_CAO_GIUA_KY_MOC_THU.docx"
ASSETS.mkdir(parents=True, exist_ok=True)

GREEN = "174C3C"
CORAL = "C75B4A"
GOLD = "B88A44"
INK = "1E2522"
MUTED = "66736D"
LIGHT = "F3F6F3"
WHITE = "FFFFFF"
BORDER = "D7DED9"


def color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run(run, size=10.5, bold=False, italic=False, value=INK, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color(value)
    return run


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def make_architecture(output: Path) -> None:
    image = Image.new("RGB", (1600, 820), "#f7f8f6")
    draw = ImageDraw.Draw(image)
    bold = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 38)
    regular = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 27)
    draw.text((70, 48), "Kiến trúc Website sách Mộc Thư", font=bold, fill="#174c3c")
    boxes = [
        (80, 210, 430, 530, "TRÌNH DUYỆT", "HTML5 / CSS3\nBootstrap / JavaScript\nResponsive UI", "#e7f0eb"),
        (625, 210, 975, 530, "PHP 8.3 + APACHE", "Route PHP\nSession / Cookie\nPDO / Validation", "#fff2ed"),
        (1170, 210, 1520, 530, "MYSQL 8.4", "8 bảng dữ liệu\nPK / FK\n1-1 / 1-n / n-n", "#fff8e9"),
    ]
    for x1, y1, x2, y2, heading, body, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=14, fill=fill, outline="#b8c4bd", width=3)
        draw.text((x1 + 30, y1 + 36), heading, font=bold, fill="#174c3c")
        draw.multiline_text((x1 + 30, y1 + 120), body, font=regular, fill="#1e2522", spacing=18)
    for x1, x2 in ((430, 625), (975, 1170)):
        draw.line((x1 + 20, 365, x2 - 32, 365), fill="#c75b4a", width=8)
        draw.polygon([(x2 - 52, 347), (x2 - 20, 365), (x2 - 52, 383)], fill="#c75b4a")
    draw.text((110, 650), "HTTP request/response", font=regular, fill="#66736d")
    draw.text((650, 650), "PDO prepared statement", font=regular, fill="#66736d")
    draw.text((260, 735), "Docker Compose cố định môi trường để chạy giống nhau trên Windows, macOS và Linux", font=regular, fill="#174c3c")
    image.save(output)


def make_code_image(source: Path, start: int, count: int, output: Path, title: str) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    snippet = lines[start - 1:start - 1 + count]
    font = ImageFont.truetype("/System/Library/Fonts/SFNSMono.ttf", 22)
    title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 24)
    width = 1700
    height = 105 + len(snippet) * 34 + 32
    image = Image.new("RGB", (width, height), "#111a17")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((15, 15, width - 15, height - 15), radius=16, fill="#111a17", outline="#375047", width=2)
    draw.text((45, 33), title, font=title_font, fill="#f3f6f3")
    keywords = re.compile(r"\b(function|if|else|return|SELECT|INSERT|UPDATE|DELETE|CREATE|TABLE|FOREIGN|KEY|REFERENCES|try|catch|const|await|fetch)\b")
    y = 88
    for number, line in enumerate(snippet, start=start):
        draw.text((35, y), f"{number:>3}", font=font, fill="#71867d")
        x = 118
        for part in keywords.split(line.expandtabs(2)):
            draw.text((x, y), part, font=font, fill="#f08b79" if keywords.fullmatch(part) else "#dce8e2")
            x += draw.textlength(part, font=font)
        y += 34
    image.save(output)


ARCH = ASSETS / "architecture.png"
CODE_DB = ASSETS / "code-database.png"
CODE_LOGIN = ASSETS / "code-login.png"
CODE_AJAX = ASSETS / "code-ajax.png"
make_architecture(ARCH)
make_code_image(ROOT / "database/schema.sql", 1, 20, CODE_DB, "database/schema.sql - users, profiles và khóa ngoại")
make_code_image(ROOT / "login.php", 1, 24, CODE_LOGIN, "login.php - xác thực và tạo session")
make_code_image(ROOT / "public/assets/js/app.js", 29, 31, CODE_AJAX, "app.js - tìm kiếm sách bằng Ajax")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.font.color.rgb = color(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, before, after, value in (
    ("Heading 1", 16, 16, 8, GREEN),
    ("Heading 2", 13, 12, 6, GREEN),
    ("Heading 3", 12, 8, 4, CORAL),
):
    style = doc.styles[name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color(value)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for list_style in ("List Bullet", "List Number"):
    style = doc.styles[list_style]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.1

header = section.header.paragraphs[0]
set_run(header.add_run("MỘC THƯ  |  BÁO CÁO ĐỒ ÁN WEBSITE SÁCH"), size=8.5, bold=True, value=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Trang "), size=8.5, value=MUTED)
add_page_field(footer)


def page_break() -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def page_title(kicker: str, title: str, intro: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(kicker.upper()), size=8.5, bold=True, value=CORAL)
    h = doc.add_paragraph(style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    set_run(h.add_run(title), size=22, bold=True, value=GREEN)
    if intro:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        set_run(p.add_run(intro), size=11, italic=True, value=MUTED)


def para(text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), bold=True, value=GREEN)
        set_run(p.add_run(text[len(bold_prefix):]))
    else:
        set_run(p.add_run(text))


def bullets(items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run(p.add_run(item))


def add_table(headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, heading in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], GREEN)
        p = table.rows[0].cells[index].paragraphs[0]
        set_run(p.add_run(heading), size=9.5, bold=True, value=WHITE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            set_run(p.add_run(value), size=9)
            if len(table.rows) % 2 == 0:
                set_cell_shading(cells[index], LIGHT)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_picture(path: Path, width=6.2, caption: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption or path.stem.replace("-", " "))
    picture._inline.docPr.set("title", path.stem)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(8)
        set_run(c.add_run(caption), size=8.5, italic=True, value=MUTED)


# 1 - Bìa
doc.add_paragraph().paragraph_format.space_after = Pt(32)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("TRƯỜNG: ........................................................"), size=11, bold=True, value=MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("KHOA: ..........................................................."), size=11, value=MUTED)
doc.add_paragraph().paragraph_format.space_after = Pt(48)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("BÁO CÁO ĐỒ ÁN WEBSITE"), size=14, bold=True, value=CORAL)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_run(p.add_run("MỘC THƯ"), size=34, bold=True, value=GREEN)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(p.add_run("Website sách sử dụng PHP 8.3 và MySQL 8.4"), size=15, value=INK)
add_picture(ROOT / "public/assets/cover-literary.png", width=2.25)
for label in ("Sinh viên: ....................................................", "Mã sinh viên: ...............................................", "Lớp: ..............................................................", "Giảng viên: ...................................................."):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(p.add_run(label), size=11, value=MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(22)
set_run(p.add_run("TP. Hồ Chí Minh, tháng 08 năm 2026"), size=10.5, italic=True, value=MUTED)
page_break()

# 2
page_title("Lời mở đầu", "Lời cảm ơn và cam đoan")
para("Em xin chân thành cảm ơn giảng viên đã cung cấp kiến thức nền tảng về HTML, CSS, JavaScript, PHP và MySQL, đồng thời định hướng phương pháp phân tích, thiết kế và kiểm thử một website hoàn chỉnh.")
para("Báo cáo này mô tả quá trình xây dựng Website sách Mộc Thư. Sản phẩm được triển khai từ yêu cầu đề bài, có mã nguồn, cơ sở dữ liệu, dữ liệu mẫu, tài liệu hướng dẫn và môi trường Docker để chạy thống nhất trên nhiều máy.")
para("Em cam đoan nội dung báo cáo phản ánh đúng sản phẩm đã triển khai. Những hình ảnh giao diện và đoạn mã minh họa được lấy từ chính dự án. Các công nghệ mã nguồn mở được sử dụng đúng mục đích học tập.")
doc.add_heading("Thông tin cần hoàn thiện trước khi nộp", level=2)
bullets(["Điền tên trường, khoa, sinh viên, mã sinh viên, lớp và giảng viên.", "Kiểm tra lại thời gian, tên học phần và quy định trình bày của khoa.", "Đính kèm link kho mã nguồn nếu giảng viên yêu cầu nộp trực tuyến."])
page_break()

# 3
page_title("Mục lục", "Cấu trúc báo cáo", "Báo cáo được tổ chức theo luồng từ yêu cầu đến thiết kế, triển khai, kiểm thử và bàn giao.")
toc = [
    ("1", "Tổng quan đề tài", "4"), ("2", "Phân tích yêu cầu", "5"), ("3", "Công nghệ sử dụng", "6"),
    ("4", "Kiến trúc hệ thống", "7"), ("5", "Thiết kế cơ sở dữ liệu", "8-10"),
    ("6", "Thiết kế giao diện công khai", "11-13"), ("7", "Xác thực và bảo mật", "14-15"),
    ("8", "Quản trị và CRUD", "16-18"), ("9", "JavaScript và CSS", "19-20"),
    ("10", "Kiểm thử", "21-22"), ("11", "Triển khai Docker", "23"),
    ("12", "Đoạn mã chính", "24"), ("13", "Phân tích chi tiết triển khai", "25-41"),
    ("14", "Kết luận và phụ lục", "42-43"),
]
add_table(["STT", "Nội dung", "Trang"], [list(row) for row in toc], [900, 6960, 1500])
page_break()

# 4
page_title("Chương 1", "Tổng quan đề tài", "Mộc Thư là website giới thiệu và quản lý sách theo hướng tối giản, ấm áp và dễ sử dụng.")
doc.add_heading("1.1 Lý do chọn đề tài", level=2)
para("Sách là loại sản phẩm có nhiều thuộc tính liên quan như danh mục, tác giả, nhà xuất bản, ISBN, giá và tồn kho. Đề tài phù hợp để thực hành giao diện responsive, xử lý form, quan hệ cơ sở dữ liệu và phân quyền quản trị.")
doc.add_heading("1.2 Mục tiêu", level=2)
bullets(["Xây dựng đầy đủ các trang công khai theo yêu cầu.", "Áp dụng HTML semantic, CSS responsive, Bootstrap và JavaScript tương tác.", "Triển khai đăng ký, đăng nhập, phân quyền và CRUD bằng PHP.", "Thiết kế tối thiểu 6 bảng MySQL có khóa chính, khóa ngoại và quan hệ.", "Đóng gói môi trường để chuyển máy và chạy bằng một lệnh."])
doc.add_heading("1.3 Phạm vi", level=2)
para("Phiên bản hiện tại tập trung vào giới thiệu danh mục sách, nội dung tin tức, tiếp nhận liên hệ và quản trị dữ liệu. Thanh toán trực tuyến và vận chuyển thực tế nằm ngoài phạm vi của học phần.")
page_break()

# 5
page_title("Chương 2", "Phân tích yêu cầu", "Mỗi nhóm yêu cầu được ánh xạ trực tiếp sang chức năng hoặc thành phần trong source.")
add_table(
    ["Nhóm", "Yêu cầu", "Hiện thực"],
    [
        ["Giao diện", "8 trang tối thiểu, responsive", "Trang chủ, giới thiệu, sách, tin tức, liên hệ, đăng nhập, đăng ký, admin"],
        ["HTML", "Heading, form, table, iframe, semantic", "header/nav/main/section/article/footer; form; bảng admin; Google Maps"],
        ["CSS", "Flex, grid, animation, media query", "style.css có grid sản phẩm, transition, reveal, breakpoint mobile"],
        ["JavaScript", "Validation, DOM, event, slider, Ajax", "Bootstrap validation, DOM toast, carousel, fetch tìm sách"],
        ["PHP", "CRUD, upload, session, cookie, phân quyền", "PDO, admin sách, upload 3 MB, role admin/customer"],
        ["Database", "Từ 6 bảng, PK/FK, quan hệ", "8 bảng chuẩn hóa, quan hệ 1-1, 1-n và n-n"],
    ],
    [1500, 3000, 4860],
)
doc.add_heading("Tiêu chí nghiệm thu", level=2)
para("Các trang phải trả về HTTP 200, không có lỗi cú pháp PHP, dữ liệu mẫu được tạo tự động, tài khoản admin đăng nhập được, CRUD và upload hoạt động, giao diện không tràn ngang ở desktop và mobile.")
page_break()

# 6
page_title("Chương 3", "Công nghệ sử dụng")
add_table(
    ["Công nghệ", "Phiên bản", "Vai trò"],
    [
        ["HTML5", "Tiêu chuẩn trình duyệt", "Cấu trúc semantic và nội dung trang"],
        ["CSS3", "Tiêu chuẩn trình duyệt", "Thiết kế, responsive, animation, transition"],
        ["Bootstrap", "5.3.3", "Grid, navbar, carousel, form và tiện ích responsive"],
        ["JavaScript", "ES2020+", "DOM, event, Ajax, validation và hiệu ứng"],
        ["PHP", "8.3", "Nghiệp vụ, xác thực, session, upload, CRUD"],
        ["MySQL", "8.4", "Lưu trữ dữ liệu quan hệ"],
        ["Apache", "2.4", "Web server cho PHP"],
        ["Docker Compose", "Compose v2", "Cố định môi trường và triển khai đa nền tảng"],
    ],
    [2100, 1600, 5660],
)
doc.add_heading("Lựa chọn kỹ thuật", level=2)
para("PDO được dùng thay cho nối chuỗi SQL để hỗ trợ prepared statement. Bootstrap được lưu cục bộ trong source, nhờ đó giao diện không phụ thuộc CDN sau khi tải Docker image.")
page_break()

# 7
page_title("Chương 4", "Kiến trúc hệ thống")
add_picture(ARCH, width=6.35, caption="Hình 4.1. Kiến trúc ba lớp đơn giản của Website Mộc Thư")
para("Trình duyệt gửi request đến Apache. Mỗi route PHP nạp các hàm dùng chung, kiểm tra session/CSRF khi cần, thực hiện truy vấn PDO và trả HTML hoặc JSON. MySQL lưu dữ liệu bền vững trong Docker volume.")
doc.add_heading("Nguyên tắc tổ chức source", level=2)
bullets(["Tách cấu hình kết nối trong config.", "Tách layout, xác thực và helper trong includes.", "Tách trang quản trị trong admin và API Ajax trong api.", "Tập trung tài nguyên tĩnh trong public/assets."])
page_break()

# 8
page_title("Chương 5", "Tổng quan cơ sở dữ liệu", "Schema gồm 8 bảng, vượt yêu cầu tối thiểu 6 bảng.")
add_table(
    ["Bảng", "Mục đích", "Khóa/quan hệ chính"],
    [
        ["users", "Tài khoản và vai trò", "PK id"],
        ["profiles", "Thông tin cá nhân", "PK/FK user_id; 1-1 users"],
        ["categories", "Danh mục sách", "PK id; 1-n books"],
        ["authors", "Tác giả", "PK id"],
        ["books", "Thông tin sách", "PK id; FK category_id"],
        ["book_authors", "Liên kết sách - tác giả", "PK kép; n-n"],
        ["news", "Bài viết tin tức", "PK id; FK author_user_id"],
        ["contacts", "Liên hệ khách hàng", "PK id"],
    ],
    [1900, 3560, 3900],
)
para("Các cột lặp được tách thành bảng riêng. Tên danh mục và tác giả không lưu trực tiếp trong books, tránh dư thừa và sai lệch khi cập nhật.")
page_break()

# 9
page_title("Chương 5", "Chi tiết bảng dữ liệu")
add_table(
    ["Bảng", "Trường tiêu biểu", "Ràng buộc"],
    [
        ["users", "email, password_hash, role, status", "email UNIQUE; role ENUM"],
        ["profiles", "user_id, full_name, phone, address", "user_id vừa là PK vừa là FK"],
        ["books", "title, slug, isbn, price, stock, status", "slug và isbn UNIQUE; price >= 0"],
        ["book_authors", "book_id, author_id", "PK kép; xóa cascade"],
        ["news", "title, slug, excerpt, content", "slug UNIQUE; FK người viết"],
        ["contacts", "name, email, subject, message", "status ENUM; thời gian tạo"],
    ],
    [1900, 4200, 3260],
)
doc.add_heading("Dữ liệu mẫu", level=2)
para("Schema khởi tạo 2 tài khoản, 5 danh mục, 5 tác giả, 10 sách và 3 bài tin. Tài khoản admin dùng bcrypt; mật khẩu mẫu là password và chỉ phục vụ môi trường học tập.")
page_break()

# 10
page_title("Chương 5", "Quan hệ và chuẩn hóa")
doc.add_heading("Quan hệ 1-1", level=2)
para("users và profiles có quan hệ 1-1. Khóa chính profiles.user_id đồng thời là khóa ngoại đến users.id.")
doc.add_heading("Quan hệ 1-n", level=2)
para("Một category có nhiều books; một user quản trị có thể viết nhiều news. Khóa ngoại đặt ở bảng phía nhiều.")
doc.add_heading("Quan hệ n-n", level=2)
para("Một sách có thể có nhiều tác giả và một tác giả có thể viết nhiều sách. Bảng book_authors giải quyết quan hệ n-n bằng khóa chính ghép.")
doc.add_heading("Chuẩn hóa", level=2)
bullets(["1NF: mỗi ô chứa một giá trị nguyên tố.", "2NF: bảng liên kết không có thuộc tính phụ thuộc một phần vào khóa ghép.", "3NF: thông tin danh mục, tác giả và hồ sơ được tách khỏi bảng sách/tài khoản.", "Ràng buộc UNIQUE và FK bảo vệ tính toàn vẹn."])
add_picture(CODE_DB, width=6.25, caption="Hình 5.1. Trích đoạn tạo bảng users và profiles")
page_break()

# 11
page_title("Chương 6", "Trang chủ")
add_picture(DOCS / "screenshots/mocthu-01-trang-chu.png", width=6.35, caption="Hình 6.1. Trang chủ Mộc Thư trên màn hình desktop")
para("Trang chủ sử dụng Bootstrap Carousel làm slider, hình ảnh thật làm tín hiệu thị giác đầu tiên, CTA dẫn đến danh mục sách và các khối sách nổi bật/tin tức ở phía dưới.")
bullets(["Heading và paragraph tạo phân cấp nội dung.", "Ảnh có alt text hỗ trợ accessibility.", "Card sách dùng CSS Grid của Bootstrap.", "Animation reveal chỉ chạy khi phần tử đi vào viewport."])
page_break()

# 12
page_title("Chương 6", "Danh mục, tìm kiếm và phân trang")
add_picture(DOCS / "screenshots/mocthu-02-danh-muc-sach.png", width=6.35, caption="Hình 6.2. Danh mục sách, bộ lọc và lưới sản phẩm")
para("Trang sách hỗ trợ tìm theo tên, ISBN hoặc nhà xuất bản; lọc theo danh mục; và phân trang 8 sản phẩm mỗi trang. Tham số được truyền vào PDO prepared statement.")
doc.add_heading("Tìm kiếm Ajax", level=2)
para("Khi nhập từ hai ký tự, JavaScript debounce 260 ms rồi gọi api/search-books.php. Kết quả JSON được dựng thành danh sách gợi ý mà không tải lại trang.")
page_break()

# 13
page_title("Chương 6", "Giới thiệu, tin tức và liên hệ")
add_table(
    ["Trang", "Thành phần HTML", "Chức năng"],
    [
        ["Giới thiệu", "section, article, list, image", "Trình bày câu chuyện, giá trị và quy trình chọn sách"],
        ["Tin tức", "article, heading, hyperlink", "Danh sách và chi tiết bài viết"],
        ["Liên hệ", "form, label, input, textarea", "Lưu nội dung liên hệ vào MySQL"],
        ["Bản đồ", "iframe", "Nhúng Google Maps theo yêu cầu HTML"],
    ],
    [1800, 3300, 4260],
)
doc.add_heading("Validation liên hệ", level=2)
para("Trình duyệt kiểm tra required và định dạng email. Máy chủ kiểm tra lại độ dài, CSRF token và dùng prepared statement trước khi INSERT, tránh phụ thuộc hoàn toàn vào client.")
page_break()

# 14
page_title("Chương 7", "Đăng ký, đăng nhập và đăng xuất")
doc.add_heading("Đăng ký", level=2)
bullets(["Kiểm tra họ tên, email, mật khẩu tối thiểu 8 ký tự.", "Kiểm tra email trùng trước khi tạo.", "Băm mật khẩu bằng password_hash.", "Tạo users và profiles trong transaction."])
doc.add_heading("Đăng nhập", level=2)
bullets(["Truy vấn theo email.", "Xác minh password_verify.", "Tạo session chứa id, email, role và full_name.", "Cookie chỉ ghi nhớ email khi người dùng chọn."])
doc.add_heading("Đăng xuất", level=2)
para("Session được xóa và tái tạo thông báo flash, sau đó chuyển về trang chủ.")
add_picture(CODE_LOGIN, width=6.25, caption="Hình 7.1. Đoạn mã xử lý đăng nhập")
page_break()

# 15
page_title("Chương 7", "Phân quyền và bảo mật")
add_table(
    ["Biện pháp", "Cách áp dụng", "Mục tiêu"],
    [
        ["Phân quyền", "require_admin kiểm tra role", "Chặn người thường vào admin"],
        ["CSRF", "Token session trong form POST", "Chặn request giả mạo"],
        ["SQL Injection", "PDO prepared statement", "Không ghép dữ liệu người dùng vào SQL"],
        ["XSS", "Hàm e() dùng htmlspecialchars", "Mã hóa dữ liệu trước khi render"],
        ["Mật khẩu", "Bcrypt qua password_hash", "Không lưu mật khẩu thô"],
        ["Upload", "Kiểm tra MIME, dung lượng, tên ngẫu nhiên", "Giảm nguy cơ file độc hại"],
        ["Cookie", "HttpOnly và SameSite=Lax", "Giảm truy cập script và CSRF"],
    ],
    [1900, 3900, 3560],
)
para("Đây là mức bảo vệ phù hợp với đồ án học tập. Khi triển khai thật cần bổ sung HTTPS, biến môi trường bí mật, rate limit và nhật ký audit.")
page_break()

# 16
page_title("Chương 8", "Bảng điều khiển quản trị")
add_picture(DOCS / "screenshots/mocthu-03-quan-tri.png", width=6.35, caption="Hình 8.1. Dashboard quản trị và thống kê cơ bản")
para("Dashboard thống kê tổng sách, sách đang bán, tài khoản và liên hệ mới. Bảng gần đây hỗ trợ người quản trị nắm nhanh trạng thái dữ liệu.")
doc.add_heading("Phạm vi admin", level=2)
bullets(["Dashboard thống kê.", "Danh sách sách có tìm kiếm và phân trang.", "Thêm, sửa, xóa sách.", "Xem và cập nhật trạng thái liên hệ."])
page_break()

# 17
page_title("Chương 8", "CRUD sách")
add_table(
    ["Thao tác", "Route", "Xử lý"],
    [
        ["Create", "admin/book-form.php", "INSERT books và book_authors trong transaction"],
        ["Read", "admin/books.php", "SELECT có join category, author; tìm kiếm; phân trang"],
        ["Update", "admin/book-form.php?id=...", "UPDATE books; thay liên kết tác giả"],
        ["Delete", "admin/book-delete.php", "POST + CSRF; DELETE; FK cascade bảng liên kết"],
    ],
    [1500, 3000, 4860],
)
doc.add_heading("Nguyên tắc xử lý", level=2)
bullets(["Không xóa bằng GET.", "ISBN phải duy nhất.", "Slug được tạo từ tiêu đề và hậu tố.", "Transaction bảo đảm sách và tác giả không lưu dở dang.", "Thông báo flash phản hồi kết quả thao tác."])
page_break()

# 18
page_title("Chương 8", "Upload ảnh bìa")
para("Admin có thể chọn JPG, PNG hoặc WEBP tối đa 3 MB. PHP kiểm tra mã lỗi upload, MIME thực bằng finfo, kích thước và tạo tên ngẫu nhiên trước khi move_uploaded_file.")
doc.add_heading("Cấu hình môi trường", level=2)
para("docker/php.ini đặt upload_max_filesize=4M và post_max_size=5M. Giới hạn PHP cao hơn giới hạn nghiệp vụ 3 MB để ứng dụng có thể tự hiển thị thông báo đúng.")
doc.add_heading("Kết quả kiểm thử", level=2)
bullets(["Upload ảnh PNG 2,7 MB thành công.", "File được lưu trong Docker volume book_uploads.", "Đường dẫn /uploads/... hiển thị được ở danh sách admin.", "File sai MIME hoặc lớn hơn 3 MB bị từ chối."])
para("Tên file do hệ thống tạo giúp tránh ghi đè và giảm rủi ro từ tên file người dùng.")
page_break()

# 19
page_title("Chương 9", "JavaScript: DOM, Event, Slider và Ajax")
add_picture(CODE_AJAX, width=6.25, caption="Hình 9.1. Tìm kiếm Ajax bằng fetch và debounce")
bullets(["Validation: bắt sự kiện submit và gọi checkValidity.", "DOM: tạo thông báo yêu thích bằng document.createElement.", "Event: scroll, click, input và DOMContentLoaded.", "Slider: Bootstrap Carousel tự động chuyển hero.", "Menu: navbar collapse hoạt động ở mobile.", "Hiệu ứng: IntersectionObserver thêm class khi card xuất hiện.", "Ajax: fetch API trả JSON gợi ý sách."])
page_break()

# 20
page_title("Chương 9", "CSS, responsive và animation")
add_table(
    ["Yêu cầu CSS", "Ví dụ trong dự án"],
    [
        ["Selector", "class, descendant, pseudo-class :hover và media feature"],
        ["Color/Font", "Biến màu xanh, coral, vàng; Apple system font stack"],
        ["Background/Border", "Hero image, nền ivory, đường viền nhẹ"],
        ["Margin/Padding", "section-space, container và spacing form"],
        ["Position", "sticky header, toast fixed"],
        ["Flexbox/Grid", "navbar, footer-grid, card layout"],
        ["Animation/Transition", "reveal, notice, hover transform"],
        ["Media Query", "Breakpoint 991px, 767px và prefers-reduced-motion"],
    ],
    [2600, 6760],
)
para("Kiểm thử ở 1280x800 và 390x844 cho thấy không có tràn ngang. Menu desktop chuyển thành nút hamburger ở mobile; kích thước chữ và khoảng cách được điều chỉnh theo breakpoint.")
page_break()

# 21
page_title("Chương 10", "Kế hoạch kiểm thử")
add_table(
    ["Mã", "Hạng mục", "Dữ liệu/Thao tác", "Kỳ vọng"],
    [
        ["TC01", "Route công khai", "GET 7 trang", "HTTP 200"],
        ["TC02", "Ajax", "Tìm 'Ngôi Nhà'", "Trả đúng 1 gợi ý"],
        ["TC03", "Đăng nhập", "admin/password", "Vào dashboard"],
        ["TC04", "Phân quyền", "User thường vào admin", "Chuyển hướng/chặn"],
        ["TC05", "CRUD", "Thêm, sửa, xóa sách test", "Dữ liệu thay đổi đúng"],
        ["TC06", "Upload", "PNG 2,7 MB", "Lưu và hiển thị ảnh"],
        ["TC07", "Responsive", "1280px và 390px", "Không overflow"],
        ["TC08", "Cài mới", "Project Docker riêng", "8 bảng, seed đăng nhập"],
    ],
    [900, 2100, 3100, 3260],
)
doc.add_heading("Phương pháp", level=2)
para("Kết hợp kiểm tra cú pháp PHP, truy vấn MySQL, HTTP curl, kiểm thử trình duyệt thật và đọc log container.")
page_break()

# 22
page_title("Chương 10", "Kết quả kiểm thử")
add_table(
    ["Hạng mục", "Kết quả", "Ghi chú"],
    [
        ["Cú pháp PHP", "Đạt", "Tất cả file PHP không có syntax error"],
        ["Route", "Đạt", "Trang HTML 200; API JSON 200"],
        ["Schema", "Đạt", "8 bảng tạo tự động"],
        ["Seed admin", "Đạt", "password_verify trả true"],
        ["CRUD", "Đạt", "Đã thêm, sửa và xóa bản ghi test"],
        ["Upload", "Đạt", "Ảnh 2,7 MB tải thành công"],
        ["Desktop/mobile", "Đạt", "Không tràn ngang, hình ảnh tải đúng"],
        ["Log PHP/Apache", "Đạt", "Không có warning/error runtime"],
    ],
    [2900, 1600, 4860],
)
para("Bản ghi và file dùng trong kiểm thử đã được xóa sau khi hoàn tất. Dữ liệu mẫu ban đầu vẫn giữ nguyên.")
page_break()

# 23
page_title("Chương 11", "Triển khai bằng Docker")
doc.add_heading("Yêu cầu", level=2)
para("Máy chạy chỉ cần Docker Desktop. Không cần cài Node.js, PHP, Apache hoặc MySQL riêng.")
doc.add_heading("Lệnh khởi động", level=2)
p = doc.add_paragraph()
set_cell = p.add_run("docker compose up --build -d")
set_run(set_cell, size=12, bold=True, value=GREEN, font="Courier New")
doc.add_heading("Quy trình tự động", level=2)
bullets(["Build image PHP 8.3 + Apache + pdo_mysql.", "Khởi động MySQL 8.4 và chờ healthcheck.", "Nạp database/schema.sql ở lần tạo volume đầu tiên.", "Khởi động web tại http://localhost:8080.", "Giữ database và upload trong named volume."])
doc.add_heading("Khả năng chuyển máy", level=2)
para("Docker Compose cố định phiên bản runtime. Bootstrap nằm trong source. File StartRun.md hướng dẫn ngắn gọn bằng tiếng Việt. Cổng có thể đổi qua APP_PORT và MYSQL_PORT.")
page_break()

# 24
page_title("Chương 12", "Ba đoạn mã chính")
doc.add_heading("12.1 Schema và quan hệ", level=2)
para("database/schema.sql định nghĩa bảng, khóa chính, khóa ngoại, unique và dữ liệu seed.")
doc.add_heading("12.2 Đăng nhập và session", level=2)
para("login.php dùng prepared statement và password_verify, sau đó tái tạo session để giảm nguy cơ session fixation.")
doc.add_heading("12.3 Tìm kiếm Ajax", level=2)
para("app.js debounce input, gọi API JSON bằng fetch, escape dữ liệu trước khi dựng HTML và hiển thị kết quả tức thời.")
add_picture(CODE_LOGIN, width=5.8)
page_break()

# 25
page_title("Chương 13", "Đặc tả chức năng tổng quát", "Phần này mô tả chi tiết hơn các chức năng đã triển khai trong source code.")
doc.add_heading("13.1 Nhóm chức năng dành cho khách truy cập", level=2)
bullets(["Xem trang chủ với slider, sách nổi bật và tin tức mới.", "Xem câu chuyện thương hiệu, giá trị và thông tin giới thiệu.", "Duyệt danh mục sách, tìm kiếm theo từ khóa và lọc theo danh mục.", "Xem chi tiết sách với tác giả, giá, ISBN, nhà xuất bản, tồn kho và mô tả.", "Đọc danh sách tin tức và trang chi tiết tin tức.", "Gửi form liên hệ để lưu vào bảng contacts.", "Đăng ký tài khoản mới và đăng nhập bằng email/mật khẩu."])
doc.add_heading("13.2 Nhóm chức năng dành cho quản trị", level=2)
bullets(["Xem dashboard thống kê tổng quan.", "Quản lý danh sách sách với phân trang và tìm kiếm.", "Thêm sách mới kèm danh mục, tác giả, giá, trạng thái và ảnh bìa.", "Sửa sách hiện có và cập nhật lại quan hệ tác giả.", "Xóa sách bằng POST có CSRF token.", "Xem danh sách liên hệ và cập nhật trạng thái xử lý."])
doc.add_heading("13.3 Ranh giới hệ thống", level=2)
para("Hệ thống không triển khai thanh toán, đặt hàng thật hoặc tích hợp vận chuyển. Các nút như thêm vào danh sách yêu thích chỉ thể hiện tương tác giao diện demo, chưa lưu trạng thái yêu thích vào database.")
page_break()

# 26
page_title("Chương 13", "Use case khách truy cập")
add_table(
    ["Use case", "Tác nhân", "Luồng chính", "Kết quả"],
    [
        ["Xem sách", "Khách", "Mở books.php, chọn danh mục hoặc nhập từ khóa", "Danh sách sách phù hợp được hiển thị"],
        ["Xem chi tiết", "Khách", "Bấm vào card sách", "book-detail.php tải dữ liệu sách theo id"],
        ["Tìm nhanh", "Khách", "Nhập tối thiểu 2 ký tự vào ô tìm kiếm", "API trả JSON và JS hiển thị gợi ý"],
        ["Gửi liên hệ", "Khách", "Điền form và gửi POST", "Dữ liệu được lưu vào contacts"],
        ["Đăng ký", "Khách", "Nhập họ tên, email, mật khẩu", "Tạo users và profiles"],
        ["Đăng nhập", "Thành viên", "Nhập email/mật khẩu hợp lệ", "Session người dùng được tạo"],
    ],
    [1700, 1600, 3860, 2200],
)
doc.add_heading("Ngoại lệ cần xử lý", level=2)
para("Nếu dữ liệu nhập không hợp lệ, trang giữ người dùng tại form hiện tại và hiển thị thông báo lỗi bằng tiếng Việt. Nếu không tìm thấy sách hoặc tin tức theo id, hệ thống trả HTTP 404 với thông báo ngắn.")
page_break()

# 27
page_title("Chương 13", "Use case quản trị")
add_table(
    ["Use case", "Điều kiện trước", "Xử lý chính", "Bảo vệ"],
    [
        ["Vào dashboard", "Đã đăng nhập admin", "Đếm sách, tài khoản, liên hệ", "require_admin"],
        ["Thêm sách", "Admin mở form", "Validate, upload ảnh, INSERT", "CSRF + transaction"],
        ["Sửa sách", "Sách tồn tại", "Load dữ liệu cũ, UPDATE", "CSRF + unique ISBN"],
        ["Xóa sách", "Sách tồn tại", "DELETE theo id", "POST + CSRF"],
        ["Xem liên hệ", "Đã đăng nhập admin", "SELECT contacts mới nhất", "require_admin"],
        ["Cập nhật liên hệ", "Có id liên hệ", "UPDATE status", "CSRF"],
    ],
    [1700, 2200, 3360, 2100],
)
para("Điểm quan trọng của khu vực admin là không để người dùng chưa xác thực truy cập route nhạy cảm. Tất cả file trong thư mục admin đều nạp includes/auth.php và gọi require_admin trước khi render giao diện.")
page_break()

# 28
page_title("Chương 13", "Luồng request và xử lý dữ liệu")
doc.add_heading("13.4 Quy trình chung của một trang PHP", level=2)
bullets(["Nạp includes/functions.php hoặc includes/auth.php.", "Đọc tham số GET/POST và chuẩn hóa kiểu dữ liệu.", "Kiểm tra điều kiện truy cập, CSRF hoặc validation nếu cần.", "Chuẩn bị câu SQL bằng PDO prepared statement.", "Thực thi truy vấn và lấy dữ liệu dạng mảng kết hợp.", "Gán pageTitle, activePage rồi include header.", "Render HTML với dữ liệu đã escape bằng hàm e().", "Include footer và nạp JavaScript cuối trang."])
doc.add_heading("13.5 Lý do chọn route PHP theo file", level=2)
para("Với phạm vi đồ án, route theo file giúp sinh viên dễ lần theo luồng xử lý từ URL đến mã nguồn. Cách này không cần framework, giảm chi phí cài đặt, nhưng vẫn đủ để thực hành tách helper, layout, xác thực và truy vấn database.")
page_break()

# 29
page_title("Chương 13", "Mapping source code với yêu cầu đề bài")
add_table(
    ["Yêu cầu", "File/thư mục", "Ghi chú"],
    [
        ["Trang chủ", "index.php", "Hero slider, sách nổi bật, tin tức"],
        ["Danh mục sách", "books.php", "Tìm kiếm, lọc, phân trang"],
        ["Chi tiết sách", "book-detail.php", "Join books, categories, authors"],
        ["Tin tức", "news.php, news-detail.php", "Danh sách và bài viết chi tiết"],
        ["Liên hệ", "contact.php", "Form, iframe Google Maps, INSERT contacts"],
        ["Auth", "login.php, register.php, logout.php", "Session, cookie, bcrypt"],
        ["Admin", "admin/*.php", "Dashboard, CRUD sách, liên hệ"],
        ["API Ajax", "api/search-books.php", "Trả JSON cho tìm kiếm nhanh"],
        ["Cấu hình", "config/*.php", "APP_URL, DB connection"],
        ["Asset", "public/assets/*", "CSS, JS, Bootstrap, ảnh bìa"],
    ],
    [2300, 2900, 4160],
)
para("Bảng mapping này giúp người chấm mở đúng file khi cần đối chiếu từng tiêu chí kỹ thuật.")
page_break()

# 30
page_title("Chương 13", "Thiết kế giao diện và trải nghiệm người dùng")
doc.add_heading("13.6 Ngôn ngữ thị giác", level=2)
para("Mộc Thư dùng tông xanh, coral và vàng nhạt để tạo cảm giác gần gũi, phù hợp với website sách. Card sản phẩm ưu tiên ảnh bìa, tiêu đề và giá để người dùng quét nhanh.")
doc.add_heading("13.7 Điều hướng", level=2)
bullets(["Header sticky giúp chuyển trang thuận tiện.", "Menu chính chỉ giữ các mục quan trọng: Trang chủ, Giới thiệu, Sách, Tin tức, Liên hệ.", "Trạng thái activePage giúp người dùng biết đang ở khu vực nào.", "Nút đăng nhập/đăng ký đổi thành lời chào và đăng xuất sau khi có session.", "Admin có tab riêng để chuyển Dashboard, Sách và Liên hệ."])
doc.add_heading("13.8 Nguyên tắc responsive", level=2)
para("Layout dùng container và grid của Bootstrap, kết hợp media query riêng để giảm kích thước hero, chuyển lưới nhiều cột thành một cột và thu gọn khoảng cách trên màn hình nhỏ.")
page_break()

# 31
page_title("Chương 13", "Phân tích trang danh mục sách")
doc.add_heading("13.9 Truy vấn chính", level=2)
para("books.php xây dựng mảng điều kiện WHERE dựa trên từ khóa và category. Với từ khóa, hệ thống tìm trong title, isbn và publisher. Với category, hệ thống lọc theo category_id.")
doc.add_heading("13.10 Phân trang", level=2)
para("Hàm pagination tính tổng số trang, ép page vào khoảng hợp lệ và trả offset. Giá trị LIMIT cố định theo ITEMS_PER_PAGE, còn OFFSET được tính từ page hiện tại.")
doc.add_heading("13.11 Hiển thị card", level=2)
bullets(["Card dùng include includes/book-card.php để tái sử dụng ở trang chủ và trang danh mục.", "Giá ưu đãi được ưu tiên nếu sale_price tồn tại.", "Tên danh mục và tác giả lấy từ JOIN/GROUP_CONCAT.", "Link chi tiết truyền id sách qua query string."])
page_break()

# 32
page_title("Chương 13", "Phân tích trang chi tiết sách")
doc.add_heading("13.12 Dữ liệu hiển thị", level=2)
bullets(["Ảnh bìa.", "Tên sách.", "Danh mục.", "Danh sách tác giả.", "Mô tả ngắn và mô tả đầy đủ.", "Giá bán, giá gốc nếu có giảm giá.", "ISBN, nhà xuất bản, năm xuất bản và tồn kho."])
doc.add_heading("13.13 Kiểm soát dữ liệu", level=2)
para("Trang chỉ hiển thị sách có status='published'. Nếu id không tồn tại hoặc sách ở trạng thái draft, người dùng nhận thông báo không tìm thấy, tránh lộ bản nháp.")
doc.add_heading("13.14 Trải nghiệm", level=2)
para("Nút quay lại danh mục giúp người dùng không bị mất ngữ cảnh duyệt sách. Nút yêu thích dùng JavaScript demo để phản hồi tức thì mà không cần backend.")
page_break()

# 33
page_title("Chương 13", "Phân tích form và validation")
add_table(
    ["Form", "Validation client", "Validation server"],
    [
        ["Đăng ký", "required, email, minlength", "filter_var email, mật khẩu >= 8, email unique"],
        ["Đăng nhập", "required, minlength", "password_verify, status active"],
        ["Liên hệ", "required, email, minlength", "độ dài tên/chủ đề/nội dung, CSRF"],
        ["Sách admin", "required, pattern ISBN, number min/max", "kiểm tra danh mục, tác giả, giá, năm, upload"],
        ["Xóa sách", "confirm trình duyệt", "POST, id hợp lệ, CSRF"],
    ],
    [1800, 3300, 4260],
)
para("Validation client giúp người dùng sửa lỗi nhanh, nhưng validation server mới là lớp quyết định vì dữ liệu POST có thể bị gửi trực tiếp từ công cụ ngoài trình duyệt.")
page_break()

# 34
page_title("Chương 13", "Bảo mật chi tiết")
doc.add_heading("13.15 CSRF token", level=2)
para("csrf_token được lưu trong session và chèn vào form qua csrf_field. Khi nhận POST, verify_csrf so sánh bằng hash_equals để tránh timing attack đơn giản.")
doc.add_heading("13.16 Escape output", level=2)
para("Hàm e() bọc htmlspecialchars với ENT_QUOTES và UTF-8. Các giá trị từ database hoặc form được escape trước khi đưa vào HTML.")
doc.add_heading("13.17 Prepared statement", level=2)
para("Các truy vấn có dữ liệu người dùng sử dụng prepare/execute. Cách này tách câu SQL khỏi tham số, giảm nguy cơ SQL Injection.")
doc.add_heading("13.18 Session và cookie", level=2)
para("Sau khi đăng nhập thành công, session_regenerate_id(true) được gọi để giảm session fixation. Cookie remember_email chỉ lưu email, không lưu mật khẩu hoặc token đăng nhập.")
page_break()

# 35
page_title("Chương 13", "Thiết kế database nâng cao")
doc.add_heading("13.19 Chỉ mục và unique", level=2)
bullets(["users.email UNIQUE để mỗi email chỉ có một tài khoản.", "categories.slug UNIQUE hỗ trợ URL ổn định nếu mở rộng route slug.", "books.slug và books.isbn UNIQUE để tránh trùng dữ liệu thương mại.", "idx_books_title hỗ trợ tìm kiếm theo tiêu đề.", "idx_books_category hỗ trợ lọc danh mục nhanh hơn."])
doc.add_heading("13.20 Hành vi khóa ngoại", level=2)
para("profiles bị xóa cascade khi user bị xóa. book_authors bị xóa cascade khi sách hoặc tác giả bị xóa. books giữ ON DELETE RESTRICT với categories để không xóa nhầm danh mục đang có sách.")
doc.add_heading("13.21 Dữ liệu kiểu ENUM", level=2)
para("role, status và trạng thái nội dung được giới hạn bằng ENUM. Điều này làm dữ liệu nhất quán, dù khi mở rộng lớn hơn có thể tách thành bảng tham chiếu riêng.")
page_break()

# 36
page_title("Chương 13", "Quy trình CRUD sách chi tiết")
doc.add_heading("13.22 Thêm mới", level=2)
para("Admin mở form rỗng, chọn danh mục/tác giả, nhập thông tin và gửi POST. Nếu dữ liệu hợp lệ, hệ thống upload ảnh, tạo slug, bắt đầu transaction, INSERT books, lấy lastInsertId rồi INSERT book_authors.")
doc.add_heading("13.23 Cập nhật", level=2)
para("Khi có id, form tải dữ liệu hiện tại. Sau khi validate, hệ thống UPDATE books, xóa liên kết cũ trong book_authors và thêm liên kết tác giả mới. Cách này đơn giản vì phiên bản hiện tại chỉ chọn một tác giả chính.")
doc.add_heading("13.24 Xóa", level=2)
para("Nút xóa nằm trong form POST, có confirm và CSRF. Khi books bị xóa, bảng book_authors tự xóa dòng liên quan nhờ ON DELETE CASCADE.")
page_break()

# 37
page_title("Chương 13", "Quản lý liên hệ")
doc.add_heading("13.25 Ghi nhận liên hệ", level=2)
para("Người dùng nhập họ tên, email, số điện thoại, chủ đề và nội dung. Nếu hợp lệ, contact.php insert bản ghi với status mặc định là new.")
doc.add_heading("13.26 Xử lý trong admin", level=2)
para("admin/contacts.php liệt kê liên hệ theo thời gian mới nhất. Người quản trị có thể đổi trạng thái sang read hoặc replied để theo dõi quy trình phản hồi.")
doc.add_heading("13.27 Giá trị nghiệp vụ", level=2)
bullets(["Không bỏ sót góp ý của khách hàng.", "Dễ phân biệt liên hệ mới và liên hệ đã xử lý.", "Tạo nền tảng để mở rộng gửi email xác nhận hoặc ticket hỗ trợ.", "Dữ liệu liên hệ giúp đánh giá nhu cầu tìm sách của người đọc."])
page_break()

# 38
page_title("Chương 13", "Docker và khả năng tái lập môi trường")
doc.add_heading("13.28 Thành phần Docker", level=2)
bullets(["Dịch vụ web dùng PHP 8.3 Apache.", "Dịch vụ db dùng MySQL 8.4.", "Schema được mount vào docker-entrypoint-initdb.d để seed lần đầu.", "Named volume db_data giữ dữ liệu MySQL.", "Named volume book_uploads giữ ảnh upload.", "Cổng web mặc định ánh xạ ra localhost:8080."])
doc.add_heading("13.29 Lợi ích khi nộp bài", level=2)
para("Giảng viên hoặc người chấm chỉ cần Docker Desktop và lệnh docker compose up --build -d. Việc này giảm lỗi khác phiên bản PHP/MySQL giữa các máy và giúp kết quả kiểm thử dễ tái hiện.")
doc.add_heading("13.30 Rủi ro", level=2)
para("Nếu máy đã chiếm cổng 8080 hoặc 3307, cần đổi APP_PORT hoặc MYSQL_PORT. Nếu muốn seed lại database, phải chạy docker compose down -v để xóa volume cũ.")
page_break()

# 39
page_title("Chương 13", "Kiểm thử chi tiết theo lớp")
add_table(
    ["Lớp", "Cách kiểm thử", "Ví dụ"],
    [
        ["Cú pháp", "php -l từng file", "Phát hiện thiếu dấu chấm phẩy, ngoặc"],
        ["Database", "SHOW TABLES, SELECT COUNT", "Kiểm tra 8 bảng và seed"],
        ["HTTP", "curl route chính", "Xác nhận 200/302/404 đúng"],
        ["Auth", "Đăng nhập admin", "Vào dashboard, thấy nút Quản trị"],
        ["Form", "Gửi dữ liệu thiếu", "Hiện lỗi validation"],
        ["CRUD", "Thêm/sửa/xóa sách test", "Dữ liệu thay đổi đúng"],
        ["Upload", "File hợp lệ và không hợp lệ", "Chấp nhận/từ chối theo MIME và size"],
        ["Responsive", "Desktop và mobile", "Không tràn ngang, menu collapse"],
    ],
    [1600, 3200, 4560],
)
para("Các kiểm thử này đủ để chứng minh chức năng chính hoạt động trong phạm vi đồ án. Khi dự án lớn hơn nên bổ sung PHPUnit hoặc Pest cho backend và Playwright cho trình duyệt.")
page_break()

# 40
page_title("Chương 13", "Đánh giá chất lượng mã nguồn")
doc.add_heading("13.31 Điểm mạnh", level=2)
bullets(["Cấu trúc thư mục rõ ràng, dễ mở rộng.", "Dùng PDO prepared statement cho thao tác database.", "Có helper dùng chung cho URL, escape, CSRF, flash message và phân trang.", "Có Docker Compose và tài liệu chạy.", "Giao diện có ảnh thật, responsive và nhất quán màu sắc.", "Database có khóa ngoại, unique và dữ liệu mẫu đầy đủ."])
doc.add_heading("13.32 Điểm còn hạn chế", level=2)
bullets(["Chưa có router tập trung hoặc controller layer.", "Chưa có test tự động.", "Chưa có quản lý nhiều tác giả trên giao diện dù database hỗ trợ n-n.", "Chưa có module đơn hàng, giỏ hàng và thanh toán.", "Chưa có chức năng biên tập tin tức trong admin.", "Thông tin cấu hình production cần tách secret chặt chẽ hơn."])
page_break()

# 41
page_title("Chương 13", "Định hướng mở rộng hệ thống")
doc.add_heading("13.33 Mở rộng chức năng bán hàng", level=2)
bullets(["Thêm bảng carts, cart_items, orders và order_items.", "Tính tổng tiền, phí vận chuyển và trạng thái đơn hàng.", "Tích hợp thanh toán chuyển khoản hoặc cổng thanh toán thử nghiệm.", "Gửi email xác nhận sau khi đặt hàng."])
doc.add_heading("13.34 Mở rộng nội dung và quản trị", level=2)
bullets(["CRUD tin tức, danh mục và tác giả trong admin.", "Upload ảnh bài viết.", "Tìm kiếm nâng cao theo khoảng giá, năm xuất bản và tình trạng tồn kho.", "Dashboard biểu đồ theo tháng và xuất báo cáo CSV."])
doc.add_heading("13.35 Mở rộng kỹ thuật", level=2)
para("Nếu tiếp tục phát triển, có thể chuyển sang mô hình MVC hoặc framework như Laravel để quản lý route, middleware, migration và test tốt hơn. Tuy vậy, bản PHP thuần hiện tại phù hợp để chứng minh nền tảng web cơ bản.")
page_break()

# 42
page_title("Chương 14", "Kết luận và hướng phát triển")
doc.add_heading("Kết quả đạt được", level=2)
para("Dự án đáp ứng đầy đủ nhóm yêu cầu giao diện, HTML, CSS, JavaScript, PHP và cơ sở dữ liệu. Website có 8 bảng chuẩn hóa, các trang công khai, xác thực, phân quyền, CRUD, upload, tìm kiếm, phân trang, thống kê và môi trường chạy đa nền tảng.")
doc.add_heading("Hạn chế", level=2)
bullets(["Chưa có giỏ hàng, thanh toán và quản lý đơn hàng.", "Chưa gửi email xác nhận hoặc khôi phục mật khẩu.", "Chưa có trang biên tập tin tức trong admin.", "Chưa triển khai HTTPS và secret manager cho production."])
doc.add_heading("Hướng phát triển", level=2)
bullets(["Bổ sung đơn hàng và thanh toán.", "Thêm đánh giá sách và danh sách yêu thích thật.", "Thêm dashboard biểu đồ và xuất báo cáo.", "Viết automated test và CI.", "Triển khai lên máy chủ có HTTPS và backup MySQL."])
page_break()

# 43
page_title("Phụ lục", "Tài liệu bàn giao và tham khảo")
doc.add_heading("A. Tài liệu trong source", level=2)
bullets(["README.md: tổng quan chức năng và cấu trúc.", "StartRun.md: cách chạy bằng Docker bằng tiếng Việt.", "database/schema.sql: thiết kế 8 bảng và dữ liệu mẫu.", "docs/screenshots/: ảnh kết quả chạy.", "docs/BAO_CAO_GIUA_KY_MOC_THU.docx: báo cáo này."])
doc.add_heading("B. Tài khoản mẫu", level=2)
para("Quản trị: admin@mocthu.vn / password. Người đọc: reader@mocthu.vn / password.")
doc.add_heading("C. Tài liệu kỹ thuật tham khảo", level=2)
bullets(["PHP Manual: PDO, session, password hashing và file upload.", "MySQL 8.4 Reference Manual: InnoDB, foreign key và constraint.", "MDN Web Docs: HTML semantic, CSS, DOM, Fetch API.", "Bootstrap 5.3 Documentation: Navbar, Carousel, Grid và Forms.", "Docker Documentation: Compose, volumes và healthcheck."])
doc.add_heading("D. Checklist trước khi nộp", level=2)
bullets(["Điền thông tin sinh viên trên bìa.", "Chạy docker compose up --build -d.", "Mở trang chủ và đăng nhập admin.", "Đính kèm link code hoặc file Moc-Thu-Source.zip.", "Kiểm tra yêu cầu đặt tên file của giảng viên."])

doc.core_properties.title = "Báo cáo đồ án Website sách Mộc Thư"
doc.core_properties.subject = "PHP 8.3, MySQL 8.4, HTML5, CSS3, JavaScript"
doc.core_properties.author = "Sinh viên"
doc.core_properties.keywords = "Mộc Thư, website sách, PHP, MySQL, Docker"
doc.save(OUTPUT)
print(OUTPUT)
